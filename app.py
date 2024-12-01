import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import threading
import os
from mistralai import Mistral
import pythoncom
import win32com.client
from datetime import datetime
import json
import re
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)

# Initialize Mistral AI API
MISTRAL_API_KEY = "API_KEY"
client = Mistral(api_key=MISTRAL_API_KEY)

class DocumentEventSystem:
    def __init__(self):
        self.events = {}
    
    def subscribe(self, event_name, callback):
        if event_name not in self.events:
            self.events[event_name] = []
        self.events[event_name].append(callback)
    
    def emit(self, event_name, data):
        if event_name in self.events:
            for callback in self.events[event_name]:
                try:
                    callback(data)
                except Exception as e:
                    logging.error(f"Error in event callback: {e}")

class DocumentEditor:
    def __init__(self):
        self.word_app = None
        self.active_doc = None
        self.changes_stack = []
        self.undo_stack = []
        self.content_cache = {}
        self.doc_lock = threading.RLock()
        
    def initialize_word(self):
        try:
            self.word_app = win32com.client.Dispatch("Word.Application")
            self.word_app.Visible = False
        except Exception as e:
            logging.error(f"Failed to initialize Word: {e}")
            raise
    
    def open_document(self, path: str) -> bool:
        if self.active_doc:
            self.close_document()
        try:
            self.active_doc = self.word_app.Documents.Open(path)
            self.active_doc.TrackRevisions = True
            return True
        except Exception as e:
            logging.error(f"Error opening document: {e}")
            return False
    
    def close_document(self):
        if self.active_doc:
            self.active_doc.Close(SaveChanges=True)
            self.active_doc = None
    
    def apply_changes(self, changes: list) -> bool:
        try:
            with self.doc_lock:
                if not self.active_doc:
                    logging.error("No active document to apply changes.")
                    return False
                
                logging.info(f"Applying {len(changes)} changes to the document.")
                
                for change in changes:
                    change_type = change.get('type')
                    if change_type == 'insert':
                        position = change.get('position', 0)
                        text = change.get('text', '')
                        rng = self.active_doc.Range(position, position)
                        rng.InsertAfter(text)
                        formatting = change.get('formatting', {})
                        if formatting:
                            self._apply_formatting(rng, formatting)
                    elif change_type == 'delete':
                        start = change.get('start', 0)
                        end = change.get('end', start)
                        rng = self.active_doc.Range(start, end)
                        rng.Delete()
                    elif change_type == 'replace':
                        start = change.get('start', 0)
                        end = change.get('end', start)
                        text = change.get('text', '')
                        rng = self.active_doc.Range(start, end)
                        rng.Text = text
                        formatting = change.get('formatting', {})
                        if formatting:
                            self._apply_formatting(rng, formatting)
                    elif change_type == 'format':
                        start = change.get('start', 0)
                        end = change.get('end', start)
                        formatting = change.get('formatting', {})
                        rng = self.active_doc.Range(start, end)
                        self._apply_formatting(rng, formatting)
                    else:
                        logging.warning(f"Unknown change type: {change_type}")
                
                self.active_doc.Save()
                self.changes_stack.append(changes)
                return True
        except Exception as e:
            logging.error(f"Unexpected error applying changes: {e}")
            return False
    
    def _apply_formatting(self, rng, formatting: dict):
        try:
            if formatting.get('bold') is not None:
                rng.Font.Bold = formatting['bold']
            if formatting.get('italic') is not None:
                rng.Font.Italic = formatting['italic']
            if formatting.get('size') is not None:
                rng.Font.Size = formatting['size']
            if formatting.get('color') is not None:
                rng.Font.Color = formatting['color']
        except Exception as e:
            logging.error(f"Error applying formatting: {e}")
    
    def undo(self) -> bool:
        if not self.changes_stack:
            return False
        changes = self.changes_stack.pop()
        self.undo_stack.append(changes)
        return True
    
    def redo(self) -> bool:
        if not self.undo_stack:
            return False
        changes = self.undo_stack.pop()
        self.apply_changes(changes)
        return True

class DocumentAnalyzer:
    def __init__(self):
        self.structure_cache = {}
        self.statistics = {}
        
    def analyze_document(self, doc: Document) -> dict:
        stats = {
            'word_count': self._count_words(doc),
            'paragraph_count': len(doc.paragraphs),
            'headings': self._analyze_headings(doc),
            'formatting': self._analyze_formatting(doc),
            'language_stats': self._analyze_language(doc)
        }
        self.statistics = stats
        return stats
        
    def _count_words(self, doc: Document) -> int:
        return sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)
        
    def _analyze_headings(self, doc: Document) -> dict:
        headings = {}
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name.split()[-1]
                headings[paragraph.text] = level
        return headings
        
    def _analyze_formatting(self, doc: Document) -> dict:
        formatting = {
            'bold_count': 0,
            'italic_count': 0,
            'underline_count': 0
        }
        return formatting
        
    def _analyze_language(self, doc: Document) -> dict:
        text = ' '.join(p.text for p in doc.paragraphs)
        return {
            'sentence_count': len(re.findall(r'[.!?]+', text)),
            'avg_word_length': sum(len(word) for word in text.split()) / len(text.split()) if text else 0
        }

class DocumentChatApp:
    def __init__(self):
        self.root = ctk.CTk()
        self.root.title("Advanced Document Chat Assistant")
        self.root.geometry("1400x900")
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")
        self.document_editor = DocumentEditor()
        self.document_analyzer = DocumentAnalyzer()
        self.current_document = None
        self.document_path = None
        self.chat_history = []
        self.event_system = DocumentEventSystem()
        self.setup_ui()
        self.setup_keyboard_shortcuts()
    
    def setup_ui(self):
        self.main_container = ctk.CTkFrame(self.root)
        self.main_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        self.setup_sidebar()
        self.setup_main_content()
    
    def setup_sidebar(self):
        sidebar = ctk.CTkFrame(self.main_container, width=200)
        sidebar.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 10))
        self.browse_button = ctk.CTkButton(sidebar, text="Open Document", command=self.browse_file)
        self.browse_button.pack(pady=5, padx=5, fill=tk.X)
        operations_frame = ctk.CTkFrame(sidebar)
        operations_frame.pack(fill=tk.X, pady=5, padx=5)
        ctk.CTkButton(operations_frame, text="Analyze Document", command=self.analyze_document).pack(fill=tk.X, pady=2)
        ctk.CTkButton(operations_frame, text="Export Chat", command=self.export_chat).pack(fill=tk.X, pady=2)
        self.stats_display = ctk.CTkTextbox(sidebar, height=200, wrap=tk.WORD)
        self.stats_display.pack(fill=tk.X, pady=5, padx=5)
        actions_frame = ctk.CTkFrame(sidebar)
        actions_frame.pack(fill=tk.X, pady=5, padx=5)
        ctk.CTkButton(actions_frame, text="↩ Undo", command=self.undo_change).pack(side=tk.LEFT, expand=True, padx=2)
        ctk.CTkButton(actions_frame, text="↪ Redo", command=self.redo_change).pack(side=tk.LEFT, expand=True, padx=2)
    
    def setup_main_content(self):
        self.chat_frame = ctk.CTkFrame(self.main_container)
        self.chat_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.chat_display = ctk.CTkTextbox(self.chat_frame, wrap=tk.WORD, font=("Arial", 12))
        self.chat_display.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        input_frame = ctk.CTkFrame(self.chat_frame)
        input_frame.pack(fill=tk.X, pady=5, padx=5)
        self.input_field = ctk.CTkTextbox(input_frame, height=100, font=("Arial", 12))
        self.input_field.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))
        send_button = ctk.CTkButton(input_frame, text="Send", command=self.send_message)
        send_button.pack(side=tk.RIGHT)
    
    def setup_keyboard_shortcuts(self):
        self.root.bind('<Control-z>', lambda e: self.undo_change())
        self.root.bind('<Control-y>', lambda e: self.redo_change())
        self.root.bind('<Control-s>', lambda e: self.save_document())
        self.input_field.bind('<Control-Return>', lambda e: self.send_message())
    
    def browse_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")])
        if file_path:
            self.document_path = file_path
            self.load_document(file_path)
    
    def load_document(self, path):
        try:
            self.current_document = Document(path)
            self.document_editor.initialize_word()
            if self.document_editor.open_document(path):
                self.append_to_chat("System", "Document loaded successfully!")
                self.analyze_document()
            else:
                self.append_to_chat("System", "Error initializing real-time editing.")
        except Exception as e:
            self.append_to_chat("System", f"Error loading document: {str(e)}")
    
    def analyze_document(self):
        if not self.current_document:
            self.append_to_chat("System", "Please load a document first!")
            return
        stats = self.document_analyzer.analyze_document(self.current_document)
        self.update_stats_display(stats)
    
    def update_stats_display(self, stats):
        self.stats_display.configure(state=tk.NORMAL)
        self.stats_display.delete('1.0', tk.END)
        self.stats_display.insert(tk.END, "Document Statistics:\n\n")
        self.stats_display.insert(tk.END, f"Words: {stats['word_count']}\n")
        self.stats_display.insert(tk.END, f"Paragraphs: {stats['paragraph_count']}\n")
        self.stats_display.insert(tk.END, f"Sentences: {stats['language_stats']['sentence_count']}\n")
        self.stats_display.insert(tk.END, f"Avg. Word Length: {stats['language_stats']['avg_word_length']:.2f}\n")
        self.stats_display.configure(state=tk.DISABLED)
    
    def send_message(self):
        message = self.input_field.get("1.0", tk.END).strip()
        command = message  # Assuming the entire message is the command
        if not message:
            return
        if not self.current_document:
            self.append_to_chat("System", "Please load a document first!")
            return
        self.input_field.delete("1.0", tk.END)
        self.append_to_chat("You", message)
        threading.Thread(target=self._process_message_thread, args=(message, command, self._update_chat_callback)).start()
    
    def _process_message_thread(self, message, command, callback):
        try:
            context = self.prepare_context(message, command)
            response = self.get_ai_response(context)
            logging.info(f"AI Response: {response}")
            callback(response, command)
        except Exception as e:
            self.root.after(0, lambda: self.append_to_chat("System", f"Error: {str(e)}"))
    
    def _update_chat_callback(self, response, command):
        self.append_to_chat("Assistant", response['message'])
        if response['status'] == 'success':
            modifications = self.parse_modifications(response['message'], command)
            if modifications:
                self.root.after(0, lambda: self.document_editor.apply_changes(modifications))
    
    def prepare_context(self, message, command) -> dict:
        return {
            'message': message,
            'command': command,
            'document_content': self.get_document_content(),
            'chat_history': self.chat_history[-5:],
            'document_stats': self.document_analyzer.statistics
        }
    
    def get_ai_response(self, context) -> dict:
        try:
            chat_response = client.agents.complete(
                agent_id="ag:364281a7:20241130:word-agent:9c4d242f",
                messages=[{"role": "user", "content": self._format_context(context)}]
            )
            return {'message': chat_response.choices[0].message.content, 'status': 'success'}
        except Exception as e:
            return {'message': f"Error getting AI response: {e}", 'status': 'error'}
    
    def _format_context(self, context: dict) -> str:
        return f"""
        Current message: {context['message']}
        
        Document content:
        {context['document_content']}
        
        Recent chat history:
        {self._format_chat_history(context['chat_history'])}
        
        Document statistics:
        {json.dumps(context['document_stats'], indent=2)}
        """
    
    def _format_chat_history(self, history: list) -> str:
        return "\n".join([f"{msg['role']}: {msg['content']}" for msg in history])
    
    def parse_modifications(self, response: str, command: str) -> list:
        try:
            pattern = r'MODIFY_DOCUMENT\s*\n*([\s\S]*?)(?=\s*$)'
            match = re.search(pattern, response)
            
            if not match:
                logging.warning("No modification found in response.")
                return []
                
            json_str = match.group(1).strip()
            json_str = re.sub(r',(\s*[}\]])', r'\1', json_str)
            modifications = json.loads(json_str)
            
            line_number = None
            search_text = None
            if "at line" in command:
                line_number_match = re.search(r'at line (\d+)', command)
                if line_number_match:
                    line_number = int(line_number_match.group(1))
            elif "find" in command:
                search_text_match = re.search(r'find "(.*?)"', command)
                if search_text_match:
                    search_text = search_text_match.group(1)
            
            if line_number:
                content = self.get_document_content()
                lines = content.split('\n')
                if line_number <= len(lines):
                    position = sum(len(line) + 1 for line in lines[:line_number - 1])
                else:
                    position = len(content)
                for change in modifications['changes']:
                    if change['type'] == 'insert':
                        change['position'] = position
            elif search_text:
                content = self.get_document_content()
                start = content.find(search_text)
                if start != -1:
                    position = start + len(search_text)
                    for change in modifications['changes']:
                        if change['type'] == 'insert':
                            change['position'] = position
                else:
                    logging.warning(f"Text '{search_text}' not found in the document.")
            
            return modifications['changes']
        except json.JSONDecodeError as e:
            logging.error(f"JSON Decode Error: {e}")
            return []
        except Exception as e:
            logging.error(f"Error parsing modifications: {e}")
            return []
    
    def append_to_chat(self, role: str, message: str):
        timestamp = datetime.now().strftime("%H:%M")
        def update_chat():
            self.chat_display.configure(state=tk.NORMAL)
            self.chat_display.insert(tk.END, f"\n[{timestamp}] {role}:\n{message}\n")
            self.chat_display.configure(state=tk.DISABLED)
            self.chat_display.see(tk.END)
        if threading.current_thread() is threading.main_thread():
            update_chat()
        else:
            self.root.after(0, update_chat)
    
    def get_document_content(self) -> str:
        if not self.current_document:
            return ""
        return "\n".join([p.text for p in self.current_document.paragraphs])
    
    def save_document(self):
        if self.document_editor.active_doc:
            try:
                self.document_editor.active_doc.Save()
                self.append_to_chat("System", "Document saved successfully!")
            except Exception as e:
                self.append_to_chat("System", f"Error saving document: {str(e)}")
    
    def export_chat(self):
        if not self.chat_history:
            messagebox.showinfo("Export Chat", "No chat history to export!")
            return
        file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt"), ("All files", "*.*")])
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    for msg in self.chat_history:
                        f.write(f"[{msg['timestamp']}] {msg['role']}:\n{msg['content']}\n\n")
                messagebox.showinfo("Export Chat", "Chat history exported successfully!")
            except Exception as e:
                messagebox.showerror("Export Error", f"Error exporting chat: {e}")
    
    def undo_change(self):
        if self.document_editor.undo():
            self.append_to_chat("System", "Change undone successfully!")
        else:
            self.append_to_chat("System", "No changes to undo.")
    
    def redo_change(self):
        if self.document_editor.redo():
            self.append_to_chat("System", "Change redone successfully!")
        else:
            self.append_to_chat("System", "No changes to redo.")
    
    def run(self):
        pythoncom.CoInitialize()
        self.root.mainloop()
        pythoncom.CoUninitialize()
    
    def cleanup(self):
        try:
            if self.document_editor.active_doc:
                self.document_editor.close_document()
            if self.document_editor.word_app:
                self.document_editor.word_app.Quit()
        except Exception as e:
            logging.error(f"Error during cleanup: {e}")
        finally:
            pythoncom.CoUninitialize()

if __name__ == "__main__":
    app = DocumentChatApp()
    try:
        app.run()
    finally:
        app.cleanup()